"""Search past project profiles by tag or keyword.

    py -m src.profiles index <profiles folder>
    py -m src.profiles search roadway central_indiana
    py -m src.profiles list-tags

Tags come from the word lists in tags.yaml.
Re-run index whenever the profiles change.
"""

import argparse
import json
import re
import sys
from datetime import datetime, timezone
from pathlib import Path

import yaml
from docx import Document

ROOT = Path(__file__).resolve().parent.parent
INDEX_PATH = ROOT / "profiles_index.json"
TAGS_PATH = ROOT / "tags.yaml"


def load_rules() -> dict:
    return yaml.safe_load(TAGS_PATH.read_text(encoding="utf-8"))


def extract_text(docx_path: Path) -> str:
    doc = Document(docx_path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def apply_tags(text: str, folder_name: str, rules: dict) -> tuple[list[str], list[str]]:
    """Return (tags, people) found in the text or the containing folder name."""
    haystack = f"{folder_name}\n{text}".lower()
    tags = []
    for category, tag_map in rules.items():
        if not isinstance(tag_map, dict):  # people / exclude_folders lists
            continue
        for tag, phrases in tag_map.items():
            if any(phrase.lower() in haystack for phrase in phrases):
                tags.append(tag)
    people = [name for name in rules.get("people", []) if name.lower() in haystack]
    return tags, people


def build_index(profiles_dir: Path) -> list[dict]:
    rules = load_rules()
    excludes = [e.lower() for e in rules.get("exclude_folders", [])]
    entries = []
    skipped = []
    for path in sorted(profiles_dir.rglob("*")):
        if path.name.startswith("~$"):  # Word lock files
            continue
        rel_parts = path.relative_to(profiles_dir).parts
        folder_path = " / ".join(rel_parts[:-1]).lower()
        if any(ex in folder_path for ex in excludes):
            continue
        if path.suffix.lower() == ".doc":
            skipped.append(path)  # old-format .doc needs Word to convert
            continue
        if path.suffix.lower() != ".docx":
            continue
        try:
            text = extract_text(path)
        except Exception as exc:  # unreadable/corrupt file — index the rest
            print(f"  ! could not read {path.name}: {exc}")
            continue
        folder = rel_parts[0] if len(rel_parts) > 1 else ""
        tags, people = apply_tags(text, folder_path, rules)
        first_line = next((ln for ln in text.splitlines() if ln.strip()), "")
        entries.append(
            {
                "file": str(path.resolve()),
                "title": first_line[:120] or path.stem,
                "service_folder": folder,
                "tags": tags,
                "people": people,
                "preview": text[:400],
                "text": text,
                "modified": datetime.fromtimestamp(
                    path.stat().st_mtime, tz=timezone.utc
                ).strftime("%Y-%m-%d"),
            }
        )
        print(f"  {path.name}: {', '.join(tags) or '(no tags matched)'}")
    if skipped:
        print(f"\nSkipped {len(skipped)} old-format .doc file(s) — open in Word and save as .docx:")
        for path in skipped:
            print(f"  {path}")
    return entries


def search_index(terms: list[str]) -> list[tuple[int, dict]]:
    if not INDEX_PATH.exists():
        sys.exit("No index yet — run:  py -m src.profiles index <profiles_folder>")
    entries = json.loads(INDEX_PATH.read_text(encoding="utf-8"))
    # Forgive punctuation people naturally type ("roadway, bridge").
    terms = [t.strip(".,;:") for t in terms if t.strip(".,;:")]
    results = []
    for entry in entries:
        tag_set = {t.lower() for t in entry["tags"]}
        people = " ".join(entry["people"]).lower()
        text = entry["text"].lower()
        score = 0
        for term in terms:
            t = term.lower()
            if any(tag.startswith(t) for tag in tag_set):
                score += 10  # tag hits (full or partial) outrank text mentions
            elif t in people:
                score += 5
            elif t in text:
                score += len(re.findall(re.escape(t), text))
            else:
                score = 0
                break  # every term must match somewhere
        if score:
            results.append((score, entry))
    results.sort(key=lambda pair: -pair[0])
    return results


def main() -> int:
    # Windows consoles default to cp1252 and choke on characters that show
    # up in real profile titles (fancy dashes, bullets). Print what we can.
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

    parser = argparse.ArgumentParser(description="Index and search project profiles.")
    sub = parser.add_subparsers(dest="command", required=True)

    p_index = sub.add_parser("index", help="Scan a folder of .docx profiles")
    p_index.add_argument("folder", help="Path to the project profiles folder")

    p_search = sub.add_parser("search", help="Search the index")
    p_search.add_argument("terms", nargs="+", help="Tags or keywords; all must match")

    sub.add_parser("list-tags", help="Show available tags from tags.yaml")

    args = parser.parse_args()

    if args.command == "index":
        folder = Path(args.folder)
        if not folder.is_dir():
            sys.exit(f"Not a folder: {folder}")
        print(f"Indexing {folder} ...")
        entries = build_index(folder)
        INDEX_PATH.write_text(
            json.dumps(entries, indent=2, ensure_ascii=False), encoding="utf-8"
        )
        print(f"\nIndexed {len(entries)} profile(s) -> {INDEX_PATH}")

    elif args.command == "search":
        results = search_index(args.terms)
        if not results:
            print("No matches. Try  py -m src.profiles list-tags  for tag names.")
            return 1
        print(f"{len(results)} match(es):\n")
        for score, entry in results:
            print(f"  {entry['title']}")
            print(f"    tags: {', '.join(entry['tags'])}")
            if entry["people"]:
                print(f"    people: {', '.join(entry['people'])}")
            print(f"    file: {entry['file']}\n")

    elif args.command == "list-tags":
        rules = load_rules()
        for category, tag_map in rules.items():
            if category == "people":
                print(f"people: {', '.join(tag_map)}")
            elif isinstance(tag_map, dict):
                print(f"{category}: {', '.join(tag_map)}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
