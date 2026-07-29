"""Compare our LOI against a competitor's, section by section.

    py -m src.compare "2605 Item 5"                  # vs the winner
    py -m src.compare "2604 Item 4" --against 2       # vs 2nd place
    py -m src.compare "2605 Item 5" --against all     # vs every ranked competitor
    py -m src.compare "2605 Item 5" --vs "C:\\...\\draft.pdf"   # a draft that
        hasn't been filed/scored yet, benchmarked against this pursuit's
        competitors as the closest available comparison

Uses the files in reference/proposal-analysis (the marketing team's folder layout).
Questions live in questions.yaml - same ones every run on purpose.
Scoring criteria go in; actual scores stay out (the marketing team's idea, avoids bias).
Each LOI goes in as text and as an image of every page, so the graphics and
layout get judged on what's actually on the page instead of guessed at from
the text.
Everything for a pursuit lands together in output/<pursuit slug>/, as both
a .docx (open this one) and a .md (source).

For running the mistake checker and this together in one pass, see
src/review.py instead.

Needs an Anthropic API key the first time:
    setx ANTHROPIC_API_KEY "sk-ant-..."   (then open a new window)
"""

import argparse
import base64
import html
import re
import sys
from datetime import date
from pathlib import Path

import anthropic
import fitz
import yaml
from docx import Document
from docx.shared import Pt
from pypdf import PdfReader

ROOT = Path(__file__).resolve().parent.parent
DATA = ROOT / "reference" / "proposal-analysis"
OUT = ROOT / "output"
MODEL = "claude-opus-5"
# --- what a run costs -----------------------------------------------------
# Page images are the expensive part of a comparison, so these two knobs are
# the cost dial. An image bills at roughly (width x height) / 750 tokens, so
# cost goes up with AREA - halving the long edge quarters the price.
#
# Measured on 2605 Item 5 (two 12-page LOIs, 24 pages, 6 questions):
#   images off       ~$0.40 a run   - graphics questions go back to guesswork
#   1100px  (default) ~$0.65 a run
#   1568px            ~$0.95 a run  - no more readable, just dearer
# The words are sent as text as well, so these only have to be good enough to
# judge layout and exhibits by - they don't have to be readable.
SEND_PAGE_IMAGES = True
MAX_IMAGE_EDGE = 1100
JPEG_QUALITY = 75
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")

INSTRUCTIONS = """You are helping VS Engineering figure out why their INDOT \
letters of interest score mid-pack. You'll get the scoring criteria, VS's LOI, \
and a competitor's LOI - each as extracted text plus an image of every page - \
then be asked about one section at a time.

Judge like an INDOT selection committee member scoring against the criteria. \
Be specific and direct: quote short phrases from both documents as evidence, \
say which firm is stronger on that section and why, and end with what VS \
should do differently. Don't pad the assessment or hedge the conclusion.

Hold to these while you do it:

Stay professional - this report gets read inside VS by the people who wrote \
the LOI. Critique the document, not the staff named in it. "His resume shows a \
bridge replacement, not an overlay" is the finding; "he may be the wrong PM" \
is not yours to say. Skip the sneering adjectives.

Neither LOI is evidence about the world. Both are sales documents making \
claims. Where the two disagree on a fact about the site, the structures, or \
past performance, report that they disagree and tell VS to check it against \
the design files - never treat the competitor's version as the truth and VS's \
as an error. Attribute claims to the document they came from.

Separate what you can see from what you're told. Page count, layout, \
graphics, and exhibits are visible to you. Anything else - a site visit, staff \
availability, a firm's track record - is a claim the document makes, and worth \
noting as one.

Weight the advice by what an LOI can actually move. Only the rated categories \
(team qualifications, project manager, project understanding and approach) come \
out of the LOI. The previous-performance categories are scored from INDOT's own \
records, and no rewrite touches them. Say which of the two you're talking about \
so nobody rewrites a page hoping to fix a number it can't reach.

Don't open your answer by restating the section name - it's already the heading \
above whatever you write."""


def pdf_text(path: Path) -> str:
    reader = PdfReader(path)
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def pdf_page_images(path: Path) -> list[dict]:
    """Every page of an LOI as a JPEG, ready to send as image content blocks.

    Extracted text loses everything the committee actually looks at - the
    plan exhibits, the detour map, the org chart, how crowded the page is.
    Without the pages the "Graphics & Layout" question was being answered
    from stray captions in the text, which is guesswork.

    We render rather than sending the PDF itself: these LOIs run up to 15MB
    of print-resolution artwork, and three of them blow past the API's 32MB
    request limit. Re-rendered at screen resolution the same document is
    about 4MB, and nothing a scorer can see is lost.
    """
    blocks = []
    with fitz.open(path) as doc:
        for page in doc:
            scale = MAX_IMAGE_EDGE / max(page.rect.width, page.rect.height)
            pixmap = page.get_pixmap(matrix=fitz.Matrix(scale, scale))
            blocks.append({
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": "image/jpeg",
                    "data": base64.standard_b64encode(
                        pixmap.tobytes("jpeg", jpg_quality=JPEG_QUALITY)
                    ).decode("ascii"),
                },
            })
    return blocks


def parse_pursuit(pursuit: str) -> tuple[str, str]:
    """'2605 Item 5' -> ('2605', '5')"""
    m = re.fullmatch(r"(\d{4})\s*item\s*(\d+)", pursuit.strip(), re.IGNORECASE)
    if not m:
        sys.exit('Pursuit should look like: "2605 Item 5"')
    return m.group(1), str(int(m.group(2)))


def default_vs_path(rfp: str, item: str) -> Path:
    """Where our own filed LOI lives for a finished, scored pursuit."""
    return DATA / "VS Proposals" / f"RFP {rfp}-VS-Item {int(item):02d}-LOI.pdf"


def available_ranks(rfp: str, item: str) -> list[int]:
    """Which competitor ranks exist for this pursuit, e.g. [1, 2, 3]."""
    comp_dir = DATA / "Competitors" / f"{rfp} Item {item}"
    ranks = {int(m.group(1)) for p in comp_dir.glob("*_*.pdf")
             if (m := re.match(r"(\d+)_", p.name))}
    return sorted(ranks)


def out_dir_for(pursuit: str) -> Path:
    """Where a pursuit's check/comparison reports live, e.g.
    output/2605-item-5/."""
    slug = re.sub(r"\W+", "-", pursuit.lower())
    return OUT / slug


DISCIPLINES_PATH = ROOT / "pursuit_disciplines.yaml"


def winners_by_discipline() -> dict[str, list[dict]]:
    """Which pursuit's winner belongs to which discipline (roadway, bridge,
    etc. - see pursuit_disciplines.yaml). Lets the web page offer "compare
    against a winner in the same discipline" even when there's no pursuit
    with the exact same RFP/Item number."""
    if not DISCIPLINES_PATH.exists():
        return {}
    tags = yaml.safe_load(DISCIPLINES_PATH.read_text(encoding="utf-8")) or {}
    out: dict[str, list[dict]] = {}
    for pursuit, disciplines in tags.items():
        rfp, item = parse_pursuit(pursuit)
        comp_dir = DATA / "Competitors" / f"{rfp} Item {item}"
        comps = sorted(comp_dir.glob("1_*.pdf"))  # 1 = the winner
        if not comps:
            continue
        who = competitor_name(comps[0], DATA / rfp)
        for d in disciplines:
            out.setdefault(d, []).append({"pursuit": pursuit, "who": who})
    return out


def find_files(pursuit: str, against: int | None, vs_path: str | Path | None = None,
               competitor_path: str | Path | None = None) -> dict:
    """Locate the VS (or draft) LOI, competitor LOI, and criteria for e.g.
    '2605 Item 5'. vs_path overrides which LOI gets scored - use it for a
    draft that hasn't been filed under VS Proposals yet. competitor_path
    overrides which competitor LOI to compare against - use it to pick a
    specific file instead of whichever is on file for `against`'s rank.
    Criteria always comes from `pursuit`'s own folder either way, since
    that's the closest available benchmark."""
    rfp, item = parse_pursuit(pursuit)

    if vs_path:
        vs = Path(vs_path)
        if not vs.exists():
            sys.exit(f"Can't find that LOI: {vs}")
    else:
        vs = default_vs_path(rfp, item)
        if not vs.exists():
            sys.exit(f"Can't find our LOI: {vs}")

    if competitor_path:
        competitor = Path(competitor_path)
        if not competitor.exists():
            sys.exit(f"Can't find that competitor LOI: {competitor}")
    else:
        comp_dir = DATA / "Competitors" / f"{rfp} Item {item}"
        comps = sorted(comp_dir.glob(f"{against}_*.pdf"))
        if not comps:
            sys.exit(f"No competitor ranked {against} in {comp_dir}")
        competitor = comps[0]

    # Criteria/RFP doc for this item, e.g. "Advertised 2604 ITEM 4 Criteria.pdf"
    # or "Advertised 2605_Item 5.pdf" - naming varies, so match loosely.
    rfp_dir = DATA / rfp
    candidates = [
        p for p in rfp_dir.glob("*.pdf")
        if re.search(rf"item\s*_?\s*0?{item}\b", p.stem, re.IGNORECASE)
        and "scoresheet" not in p.stem.replace(" ", "").lower()
    ]
    candidates.sort(key=lambda p: "criteria" not in p.stem.lower())
    if not candidates:
        sys.exit(f"No criteria/RFP doc for item {item} in {rfp_dir}")

    return {
        "vs": vs,
        "competitor": competitor,
        "criteria": candidates[0],
        "who": competitor_name(competitor, rfp_dir),
    }


def competitor_name(path: Path, rfp_dir: Path) -> str:
    """Get the firm name from the score sheet filenames, e.g.
    '2605 Item 5_FinalScoreSheet_PARSONS.pdf' -> PARSONS."""
    key = re.sub(r"[^a-z0-9]", "", path.stem.lower())
    for sheet in rfp_dir.glob("*FinalScoreSheet_*.pdf"):
        firm = sheet.stem.rsplit("_", 1)[-1]
        norm = re.sub(r"[^a-z0-9]", "", firm.lower())
        if norm != "vs" and norm in key:
            return firm.title() if firm.isupper() else firm
    return re.sub(r"^\d+_", "", path.stem)  # fall back to the filename


# --- Turns a markdown report into a Word doc or a web page. Both readers
# share one parser so the two outputs always match. Only handles what
# Claude's answers actually use: #/## headings, **bold**, and lists.

_HEADING_RE = re.compile(r"(#{1,6})\s+(.+)")
_ITEM_RE = re.compile(r"(\d+\.|[-*])\s+(.*)")


def _parse_blocks(md_text: str):
    """Split a report into (kind, content) pieces: headings, paragraphs,
    numbered lists, bulleted lists. Blocks are separated by a blank line."""
    for block in re.split(r"\n\s*\n", md_text.strip()):
        block = block.strip()
        if block:
            yield from _parse_block(block)


def _parse_block(block: str):
    """One blank-line-separated block often holds several things at once -
    Claude writes a heading, a lead-in line, and a list with no blank lines
    between them. Walk the lines and emit each run on its own, or the lead-in
    gets glued onto the list below it and the whole thing renders as one
    run-on paragraph (which is exactly what used to happen)."""
    out = []
    para: list[str] = []
    items: list[str] = []
    list_kind = None

    def flush():
        if para:
            out.append(("p", re.sub(r"\s+", " ", " ".join(para)).strip()))
            para.clear()
        if items:
            out.append((list_kind, list(items)))
            items.clear()

    for line in block.splitlines():
        line = line.strip()
        if not line:
            continue

        heading = _HEADING_RE.fullmatch(line)
        if heading:
            flush()
            level = len(heading.group(1))
            out.append(("h0" if level == 1 else "h1" if level == 2 else "h2",
                        heading.group(2).strip()))
            continue

        item = _ITEM_RE.fullmatch(line)
        if item:
            kind = "list-number" if item.group(1)[0].isdigit() else "list-bullet"
            if items and kind != list_kind:
                flush()  # a bulleted list and a numbered one aren't one list
            elif para:
                flush()  # the lead-in line above belongs to itself
            list_kind = kind
            items.append(item.group(2).strip())
            continue

        if items:
            items[-1] += " " + line  # wrapped continuation of the item above
        else:
            para.append(line)

    flush()
    return out


def _add_rich_paragraph(doc, text, style=None):
    """One Word paragraph. Turns **bold** into an actual bold run."""
    p = doc.add_paragraph(style=style)
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])  # plain text before the bold bit
        p.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])  # whatever's left after the last bold bit
    return p


def markdown_to_docx(md_text: str, out_path: Path) -> Path:
    doc = Document()
    for kind, content in _parse_blocks(md_text):
        if kind == "h0":
            doc.add_heading(content, level=0)
        elif kind == "h1":
            doc.add_heading(content, level=1)
        elif kind == "h2":
            doc.add_heading(content, level=2)
        elif kind == "p":
            _add_rich_paragraph(doc, content)
        elif kind == "list-number":
            for item in content:
                _add_rich_paragraph(doc, item, style="List Number")
        elif kind == "list-bullet":
            for item in content:
                _add_rich_paragraph(doc, item, style="List Bullet")

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


def _bold_to_html(text: str) -> str:
    """Escape the text for HTML, then turn **bold** into <strong>."""
    parts = []
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            parts.append(html.escape(text[pos:m.start()]))
        parts.append(f"<strong>{html.escape(m.group(1))}</strong>")
        pos = m.end()
    parts.append(html.escape(text[pos:]))
    return "".join(parts)


def markdown_to_html(md_text: str) -> str:
    """Same report as markdown_to_docx, as an HTML snippet - so it can be
    read right on the search page instead of opening Word."""
    out = []
    for kind, content in _parse_blocks(md_text):
        if kind == "h0":
            out.append(f"<h2>{_bold_to_html(content)}</h2>")
        elif kind == "h1":
            out.append(f"<h3>{_bold_to_html(content)}</h3>")
        elif kind == "h2":
            out.append(f"<h4>{_bold_to_html(content)}</h4>")
        elif kind == "p":
            out.append(f"<p>{_bold_to_html(content)}</p>")
        elif kind == "list-number":
            items = "".join(f"<li>{_bold_to_html(i)}</li>" for i in content)
            out.append(f"<ol>{items}</ol>")
        elif kind == "list-bullet":
            items = "".join(f"<li>{_bold_to_html(i)}</li>" for i in content)
            out.append(f"<ul>{items}</ul>")
    return "\n".join(out)


def run(pursuit: str, against: int | None, vs_path: str | Path | None = None,
        competitor_path: str | Path | None = None) -> Path:
    files = find_files(pursuit, against, vs_path, competitor_path)
    who = files.pop("who")
    # A draft gets its own filename as the label, since it isn't really
    # "VS's <pursuit> LOI" - it's being benchmarked against that pursuit's
    # winner as the closest available comparison.
    us = Path(vs_path).stem if vs_path else "VS Engineering"
    print(f"Comparing {us} vs {who} (benchmark: {pursuit})")
    for label, path in files.items():
        print(f"  {label}: {path.name}")

    # Each LOI goes in twice: the extracted text, so quotes come back with the
    # exact wording, and every page as an image, so the layout and exhibits can
    # actually be judged. The criteria doc is a table of scoring rules - text is
    # all it has to give, so it doesn't get rendered.
    docs = [
        ("SCORING CRITERIA / RFP", files["criteria"], False),
        ("VS ENGINEERING LOI", files["vs"], True),
        (f"COMPETITOR LOI ({who})", files["competitor"], True),
    ]
    content = []
    for title, path, with_pages in docs:
        content.append({"type": "text", "text": f"=== {title}: {path.name} ==="})
        content.append({"type": "text", "text": pdf_text(path)})
        if with_pages and SEND_PAGE_IMAGES:
            pages = pdf_page_images(path)
            print(f"  rendered {len(pages)} page(s) of {path.name}")
            content.append({"type": "text", "text": f"--- {title}, every page as an image ---"})
            content += pages
    # Everything above is identical on all 6 questions, so cache it here and
    # let each question's text land after the breakpoint.
    content[-1]["cache_control"] = {"type": "ephemeral"}

    questions = yaml.safe_load((ROOT / "questions.yaml").read_text(encoding="utf-8"))
    client = anthropic.Anthropic()

    report = [
        f"# LOI comparison: {us} vs {who} (benchmark: {pursuit})",
        f"\nGenerated {date.today()} by src/compare.py. "
        f"Same questions every run (questions.yaml).\n",
    ]
    for q in questions:
        print(f"  asking about: {q['section']} ...")
        response = client.messages.create(
            model=MODEL,
            # Covers thinking and the answer together, so it has to be roomy -
            # too low and a section gets cut off mid-sentence.
            max_tokens=16000,
            thinking={"type": "adaptive"},
            system=INSTRUCTIONS,
            messages=[{
                "role": "user",
                "content": content + [{"type": "text", "text": q["question"]}],
            }],
        )
        answer = "".join(b.text for b in response.content if b.type == "text").strip()
        # Never let a cut-off or refused section land in the report looking
        # like a finished one.
        if response.stop_reason == "max_tokens":
            answer += ("\n\n**This section ran out of room and stops mid-thought. "
                       "Re-run it, or raise max_tokens in src/compare.py.**")
            print("    ! hit the length limit - this section is incomplete")
        elif response.stop_reason == "refusal":
            answer = ("**The model declined to answer this section.** Nothing was "
                      "written for it. The rest of the report is unaffected.")
            print("    ! declined - section left empty")
        report.append(f"\n## {q['section']}\n\n{answer}\n")

    out_dir = out_dir_for(pursuit)
    out_dir.mkdir(parents=True, exist_ok=True)
    name = f"vs_{re.sub(r'\W+', '-', who)}"
    if vs_path:
        name = f"draft-{re.sub(r'\W+', '-', Path(vs_path).stem)}_{name}"
    stem = out_dir / name
    md_text = "\n".join(report)
    stem.with_suffix(".md").write_text(md_text, encoding="utf-8")
    markdown_to_docx(md_text, stem.with_suffix(".docx"))
    return stem.with_suffix(".docx")


def main() -> int:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    parser = argparse.ArgumentParser(description="Compare our LOI against a competitor's.")
    parser.add_argument("pursuit", help='e.g. "2605 Item 5" - which pursuit\'s '
                        "competitor data to benchmark against")
    parser.add_argument("--against", default="1",
                        help='competitor rank: 1 = winner (default), 2, 3, or "all"')
    parser.add_argument("--vs", help="Path to a specific LOI to score, instead of "
                        "the one filed under reference/proposal-analysis/VS Proposals. "
                        "Use this for a draft that hasn't been submitted/scored yet.")
    parser.add_argument("--competitor", help="Path to a specific competitor LOI to "
                        "compare against, instead of whichever is on file for "
                        "--against's rank. Ignores --against when given.")
    args = parser.parse_args()

    if args.competitor:
        ranks = [None]
    elif args.against.lower() == "all":
        rfp, item = parse_pursuit(args.pursuit)
        ranks = available_ranks(rfp, item)
        if not ranks:
            sys.exit(f"No competitors found for {args.pursuit}")
    else:
        ranks = [int(args.against)]

    try:
        out_paths = [run(args.pursuit, rank, args.vs, args.competitor) for rank in ranks]
    except anthropic.AuthenticationError:
        # A key IS set, but Anthropic rejected it.
        sys.exit("API key is set but Anthropic rejected it (wrong or revoked). "
                 "Get a fresh one at console.anthropic.com, then:\n"
                 '  setx ANTHROPIC_API_KEY "sk-ant-..."\n'
                 "then open a new window.")
    except TypeError as exc:
        if "auth" in str(exc).lower():
            sys.exit("No API key set. Run:\n"
                     '  setx ANTHROPIC_API_KEY "sk-ant-..."\n'
                     "then open a new window. Keys: console.anthropic.com")
        raise
    print(f"\nDone. {len(out_paths)} report(s):")
    for p in out_paths:
        print(f"  {p}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
