"""Assemble the pre-proposal into a Word document.

Currently builds a clean generic layout with python-docx. If the meeting
reveals a strict VSE template, swap this module to fill that template
(e.g. with docxtpl placeholders) — nothing upstream changes.
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt


def build_document(company_name: str, intake: dict, sections: list[tuple[str, str]], out_path: Path) -> Path:
    doc = Document()

    doc.add_heading(f"{company_name} — Pre-Proposal", level=0)
    doc.add_paragraph(f"Project: {intake['project_name']}")
    doc.add_paragraph(f"Prepared for: {intake['client_name']}")
    doc.add_paragraph(f"Date: {date.today().strftime('%B %d, %Y')}")

    for title, body in sections:
        doc.add_heading(title, level=1)
        for block in body.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            # Preserve simple bullet lists coming out of the drafter/library
            if all(line.lstrip().startswith(("-", "*")) for line in block.splitlines()):
                for line in block.splitlines():
                    doc.add_paragraph(line.lstrip("-* ").strip(), style="List Bullet")
            else:
                doc.add_paragraph(block)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path
